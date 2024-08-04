from flask import Flask, request, send_file
import speech_recognition as sr
import os
from pathlib import Path
import openai
from docx import Document
from docx.shared import Pt 

app = Flask(__name__)

@app.route('/')
def index():
    return "Welcome to the Speech-to-Text to Word File API"

@app.route('/upload', methods=['POST'])
def upload():
    file = request.files['file']
    
    # Speech-to-text
    recognizer = sr.Recognizer()
    audio_file = sr.AudioFile(file)
    with audio_file as source:
        audio = recognizer.record(source)
    text = recognizer.recognize_google(audio)

    # Configuration and paths
    openai.organization = 'org-yRlfrdqdXMIAYGfdaIqbyL28'
    openai_api_key = "sk-proj-iXpA1QzCyeOwS9ORRxACT3BlbkFJgZm1iSBO3S8S64bGddlS"

    # Initialize the OpenAI API client
    openai.api_key = openai_api_key

    # Supported audio formats for transcription
    supported_formats = ['flac', 'm4a', 'mp3', 'mp4', 'mpeg', 'mpga', 'oga', 'ogg', 'wav', 'webm']

    # Open the most recent audio file and transcribe it
    with open(file, "rb") as f:
        transcription = openai.Audio.transcribe("whisper-1", f)  # Assuming this is the correct OpenAI API call
    print(transcription)
    
    # System prompt
    system_prompt = """
    You are a helpful assistant for a cardiology doctor. Your task is to take the text and convert the points provided into prose. Correct any spelling and grammar discrepancies, using English UK, in the transcribed text. Maintain accuracy of the transcription and use only context provided. Format the output into a medical letter under the following headings: '###Reason for Referral/Diagnosis', '###Medications', '###Clinical Review', '###Diagnostic Tests', '###Plan', and '###Actions for GP'. The 'Reason for Referral/Diagnosis' should be a numbered list. The 'Medications' should be in a sentence, capitalise the first letter of the drug name and separate them by commas. Format the 'Clinical Review' in paragraphs for readability. Always leave the 'Diagnostic Tests' blank. Do not add any address options at the beginning or any signatures at the end.
    Important not to redact the plan from the clinical review. Keep the accurate prose plan in the clinical review, and also create a list of points for the 'Plan' and 'Actions for GP'.
    Always start the 'Clinical Review' with 'It was a pleasure reviewing [patient's name] in the Arrhythmia clinic on behalf of Dr. today. [He/She] is a [age] year old patient...'
    At the end of the letter always finish with:
    '###Signature'
    Dr. Jasjit Syan
    Cardiology Registrar
    \n
    Cardiology Department: Telephone: 020 8321 5336/Email: caw-tr.westmidadmin7@nhs.net
    Appointments: 020 8321 5610 Email: caw-tr.wm-bookingenquiries@nhs.net
    \n
    Disclaimer: This document has been transcribed from dictation; we apologize for any unintentional spelling mistakes/errors due to the voice recognition software.
    """

    def generate_corrected_transcript(temperature, system_prompt, transcribed_text):
        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=0.2,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": transcribed_text
                }
            ]
        )
        return response.choices[0].message['content']

    # Example usage
    corrected_text = generate_corrected_transcript(0.2, system_prompt, transcription['text'])
    print(corrected_text)

    def set_paragraph_format(paragraph):
        """Remove space after paragraph."""
        p_format = paragraph.paragraph_format
        p_format.space_after = Pt(0)

    def add_bold_underline_heading(doc, text, level):
        """Add a bold and underlined heading."""
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.bold = True
        run.underline = True
        set_paragraph_format(heading)

    def save_to_word(corrected_text, output_path):
        doc = Document()
        lines = corrected_text.split('\n')
        for line in lines:
            if line.startswith('###'):  # Heading level 3
                add_bold_underline_heading(doc, line[3:].strip(), level=3)
            elif line.startswith('##'):  # Heading level 2
                add_bold_underline_heading(doc, line[2:].strip(), level=2)
            elif line.startswith('#'):  # Heading level 1
                add_bold_underline_heading(doc, line[1:].strip(), level=1)
            elif line.startswith('- '):  # Bullet points
                paragraph = doc.add_paragraph(line, style='ListBullet')
                set_paragraph_format(paragraph)
            else:
                paragraph = doc.add_paragraph(line)
                set_paragraph_format(paragraph)
        doc.save(output_path)

    # Save the corrected transcript to a Word document
    output_path = os.path.join("uploads", "output.docx")
    save_to_word(corrected_text, output_path)
    
    return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    if not os.path.exists('uploads'):
        os.makedirs('uploads')
    app.run(debug=True)
