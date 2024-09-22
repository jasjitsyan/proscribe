import os
from pathlib import Path
from flask import Flask, request, render_template, jsonify, make_response
from flask_cors import CORS
import openai
from docx import Document
from docx.shared import Pt
from io import BytesIO

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configuration and paths
AUDIO_DIR = Path("./audio")
OUTPUT_DIR = Path("./text")

# Ensure directories exist
AUDIO_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# OpenAI API setup
openai.organization = 'org-yRlfrdqdXMIAYGfdaIqbyL28'
openai.api_key = os.getenv("OPENAI_API_KEY")

system_prompt = """
You are a helpful assistant for a cardiology doctor. Your task is to take the text and convert the points provided into prose. Correct any spelling and grammar discrepancies...
"""

def generate_corrected_transcript(temperature, system_prompt, transcribed_text):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=temperature,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": transcribed_text}
        ]
    )
    return response.choices[0]['message']['content']

def save_to_word(corrected_text):
    doc = Document()
    lines = corrected_text.split('\n')

    for line in lines:
        if line.startswith('###'):
            doc.add_heading(line[3:].strip(), level=3)
        else:
            paragraph = doc.add_paragraph(line)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)

    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

@app.route('/')
def index():
    return render_template('upload.html')

@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    if 'audio_file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    audio_file = request.files['audio_file']
    audio_file_path = AUDIO_DIR / audio_file.filename
    audio_file.save(audio_file_path)

    try:
        with open(audio_file_path, 'rb') as f:
            transcript = openai.Audio.transcribe("whisper-1", f)
        transcribed_text = transcript['text']
    except Exception as e:
        return jsonify({'error': str(e)}), 500

    return jsonify({'transcribedText': transcribed_text})

@app.route('/download-docx', methods=['POST'])
def download_docx():
    transcribed_text = request.form['transcribed_text']
    corrected_text = generate_corrected_transcript(0.7, system_prompt, transcribed_text)
    
    doc_stream = save_to_word(corrected_text)
    response = make_response(doc_stream.read())
    response.headers.set('Content-Disposition', 'attachment', filename='medical_letter.docx')
    response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    return response

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
